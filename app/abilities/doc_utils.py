"""文档文本抽取公共能力：txt/md/docx/pdf。无需大模型。"""
import re
from fastapi import HTTPException

ARTICLE_RE = re.compile(r"第[一二三四五六七八九十百千零〇0-9]+[条款]")


def extract_text(filename: str, raw: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext in ("txt", "md"):
        for enc in ("utf-8", "gbk", "gb18030"):
            try:
                return raw.decode(enc)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="replace")
    if ext == "docx":
        from docx import Document
        import io
        doc = Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(parts)
    if ext == "pdf":
        try:
            import fitz
        except ImportError:
            import pymupdf as fitz
        import io
        doc = fitz.open(stream=raw, filetype="pdf")
        text = "".join(page.get_text() for page in doc)
        doc.close()
        return text
    raise HTTPException(400, f"暂不支持该格式：{ext}（支持 txt/md/docx/pdf）")


def split_articles(text: str):
    """按'第X条'切片；退化按段落切片。"""
    matches = list(ARTICLE_RE.finditer(text))
    if len(matches) >= 2:
        chunks = []
        for i, m in enumerate(matches):
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            body = text[m.start():end].strip()
            if body:
                chunks.append((m.group(0), body))
        return chunks
    paras = [p.strip() for p in text.splitlines() if p.strip()]
    merged, cur = [], ""
    for p in paras:
        if len(p) > 200 and cur:
            merged.append(("", cur))
            cur = ""
        cur = (cur + "\n" + p).strip()
        if len(cur) >= 100:
            merged.append(("", cur))
            cur = ""
    if cur:
        merged.append(("", cur))
    return merged or [("", text[:2000])]
